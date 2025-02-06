import csv
import json
import os
from datetime import datetime
from typing import TextIO

import docx
from whisperx.utils import ResultWriter, format_timestamp

# This dict holds the result of the transcription with grouped speakers
merged_speakers_result:  dict = {}
# This function groups the speaker lines with the same speaker into one
def merge_speakers(result: dict) -> dict:
    # only create the grouping one time per input file
    if len(merged_speakers_result) == 0:
        # add the segments list and language
        merged_speakers_result["segments"] = []
        merged_speakers_result["language"] = result["language"]

        current_speaker: str = ""
        for line in result["segments"]:
            speaker, text = extract_speaker_and_text(line)
            if speaker != current_speaker:
                # create new dict and add to segments list
                new_segment = {
                    "speaker": speaker,
                    "text": text,
                    "start": line["start"],
                    "end": line["end"]
                }
                merged_speakers_result["segments"].append(new_segment)
                current_speaker = speaker
            else:
                # add the text to the previous dict in the segments list
                merged_speakers_result["segments"][-1]["text"] = (merged_speakers_result["segments"][-1]["text"].strip()
                                                                  + " "
                                                                  + text.strip())
                # push the end time forward
                merged_speakers_result["segments"][-1]["end"] = line["end"]

    return merged_speakers_result

# This function will clear the merged speaker data. Use the function before creating data for the next file.
def reset_merge_speaker_data():
    merged_speakers_result.clear()

def get_field_names(result: dict) -> list[str]:
    # make sure that the CSV header contains the 'speaker' header even though the first line has no speaker
    # if there is no speaker in the segment, try to take the keys from the next line
    # we don't want to hard code the speaker key into a fixed position in the header list
    for i in range(len(result)):
        if 'speaker' in result["segments"][i]:
            return list(result["segments"][i].keys())
    # if none of the lines had a speaker then just return the keys from the first line
    return list(result["segments"][0].keys())

def write_error(exception: Exception, extension: str):
    print(f"An error occurred while attempting to write the {extension} format: {exception}")

class WriteCSV(ResultWriter):
    extension: str = "csv"

    def write_result(self, result: dict, file: TextIO, options: dict):
        try:
            fieldnames: list[str] = get_field_names(result)
            clean_result_for_csv_writer(fieldnames, result)
            writer = csv.DictWriter(file, fieldnames=fieldnames)
            writer.writeheader()
            writer.writerows(result["segments"])
        except Exception as e:
            write_error(e, self.extension)
            remove_output_file(file)


class WriteDOTE(ResultWriter):
    extension: str = "dote.json"

    @staticmethod
    def format_result(result: dict):
        interface = {"lines": []}
        for line in result["segments"]:
            speaker, text = extract_speaker_and_text(line)
            line_add = {
                "startTime": format_timestamp(line["start"], True),
                "endTime": format_timestamp(line["end"], True),
                "speakerDesignation": speaker.strip(),
                "text": text.strip(),
            }
            interface["lines"].append(line_add)
        return interface

    def write_result(self, result: dict, file: TextIO, options: dict):
        try:
            transcription_data = self.format_result(result)
            json.dump(transcription_data, file)
        except Exception as e:
            write_error(e, self.extension)
            remove_output_file(file)


class WriteDOCX(ResultWriter):
    extension: str = "docx"

    @staticmethod
    def format_time(
            start_time: int | float, end_time: int | float, max_time: int | float
    ) -> str:
        dt_start = datetime.utcfromtimestamp(start_time)
        dt_end = datetime.utcfromtimestamp(end_time)
        time_format = "%M:%S"
        if max_time >= 3600:
            time_format = "%H:%M:%S"
        return f"{dt_start.strftime(time_format)} - {dt_end.strftime(time_format)}"

    def write_result(self, result: dict, file: TextIO, options: dict):
        try:
            audio_filename = options.get("filename", "").stem
            document = docx.Document()
            document.add_heading(audio_filename, level=2)
            document.add_heading(options["jobname"], level=4)
            document.extended_properties.set_property("total_time", "   1")
            document.extended_properties.set_property(
                "application", "Whisper Transcription AAU extension"
            )
            document.core_properties.title = audio_filename
            document.core_properties.author = options.get("username", "")
            document.core_properties.subject = options.get("jobname", "")

            p = document.add_paragraph()
            max_time = result["segments"][-1]["end"]
            for line in result["segments"]:
                speaker, text = extract_speaker_and_text(line)
                time = p.add_run(self.format_time(line["start"], line["end"], max_time))
                time.italic = True
                p.add_run("\t")
                p.add_run(f'{speaker.strip()}\n')
                p.add_run("\t")
                p.add_run(f'{text.strip()}\n')
            document.save(file.name)
        except Exception as e:
            write_error(e, self.extension)
            remove_output_file(file)


def extract_speaker_and_text(line):
    # if the algorithm was unable to determine the speaker, the line object will not have a speaker object
    if "speaker" in line:
        speaker = line["speaker"]
    else:
        speaker = "Undetermined speaker"
    # Also guard against missing texts if the algorithm was unable to transcribe the segment
    if "text" in line:
        text = line["text"]
    else:
        text = "Undetermined text"

    return speaker, text

def remove_output_file(file: TextIO):
    try:
        file_path = file.name
        file.close()
        os.remove(file_path)
    except Exception as e:
        print(f"An error occurred while attempting to delete the output file:  {file_path}, exception: {e}")

# This method removes data entries in the result dictionary that are not part of the header set
# see https://github.com/aau-claaudia/transcriber/issues/20
def clean_result_for_csv_writer(headers: list[str], result: dict) -> dict:
    for line in result["segments"]:
        # line is a dictionary, and we want to remove entries that have a key name which is not in the headers list
        # for this line create a set of keys to remove
        keys_to_remove = [key for key in line.keys() if key not in headers]
        for key in keys_to_remove:
            # remove the entry
            del line[key]
    return result